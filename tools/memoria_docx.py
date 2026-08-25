#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Genera la Memoria Descriptiva de SocialKids en formato .docx editable.

Sigue la estructura de la plantilla de referencia:
    1. Nombre de la aplicacion
    2. Nombre corto
    3. Version
    4. Sobre la aplicacion
       4.1 Estructura del proyecto
       4.2 Codigos principales de la aplicacion
       4.3 Relacion general de modulos y archivos del codigo fuente
       4.4 Flujo general de funcionamiento
    5. Observacion final

Los fragmentos de codigo se leen del propio repositorio, de modo que el
documento no puede quedar desincronizado con el codigo real.

Uso:
    python tools/memoria_docx.py [salida.docx]
"""
import io
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

TURQUESA = RGBColor(0x0E, 0x8F, 0x92)
CORAL = RGBColor(0xD9, 0x4A, 0x3A)
NOCHE = RGBColor(0x1B, 0x2B, 0x3F)
GRIS = RGBColor(0x55, 0x62, 0x72)


# --------------------------------------------------------------------- utiles

def leer(ruta, desde=None, hasta=None, sin_imports=False):
    """Devuelve el contenido real de un archivo del repositorio."""
    completo = io.open(os.path.join(RAIZ, ruta), encoding="utf-8").read()
    lineas = completo.split("\n")
    if desde is not None:
        lineas = lineas[desde - 1: hasta if hasta else len(lineas)]
    if sin_imports:
        lineas = [l for l in lineas if not l.startswith("import ")]
    # Quita lineas en blanco consecutivas al principio y al final
    while lineas and not lineas[0].strip():
        lineas.pop(0)
    while lineas and not lineas[-1].strip():
        lineas.pop()
    return "\n".join(lineas)


def borde_inferior(parrafo, color="12B5B0", grosor=12):
    pPr = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(grosor))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    bordes.append(bottom)
    pPr.append(bordes)


def sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


class Memoria:
    def __init__(self):
        self.doc = Document()
        self._configurar()

    # ------------------------------------------------------------ estructura

    def _configurar(self):
        seccion = self.doc.sections[0]
        seccion.page_width = Cm(21.0)
        seccion.page_height = Cm(29.7)
        seccion.left_margin = Cm(2.5)
        seccion.right_margin = Cm(2.5)
        seccion.top_margin = Cm(2.5)
        seccion.bottom_margin = Cm(2.5)

        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = NOCHE
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        # Pie de pagina con numeracion automatica
        pie = seccion.footer.paragraphs[0]
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pie.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = GRIS
        for instruccion in ("begin", "PAGE", "end"):
            elem = OxmlElement("w:fldChar" if instruccion != "PAGE" else "w:instrText")
            if instruccion == "PAGE":
                elem.set(qn("xml:space"), "preserve")
                elem.text = " PAGE "
            else:
                elem.set(qn("w:fldCharType"), instruccion)
            run._r.append(elem)

    # ---------------------------------------------------------------- bloques

    def portada(self, titulo, subtitulo, pie):
        for _ in range(6):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(34)
        r.font.color.rgb = TURQUESA

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(subtitulo)
        r.font.size = Pt(14)
        r.font.color.rgb = NOCHE

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        borde_inferior(p)

        for _ in range(8):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(pie)
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRIS

        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def h1(self, numero, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("%s   %s" % (numero, texto.upper()))
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = TURQUESA
        borde_inferior(p, "CFE2E3", 6)

    def h2(self, numero, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("%s   %s" % (numero, texto))
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = CORAL

    def h3(self, numero, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("%s   %s" % (numero, texto))
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NOCHE

    def parrafo(self, texto, cursiva=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(texto)
        r.italic = cursiva
        return p

    def punto(self, texto, nivel=0, guion="-"):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8 + nivel * 0.7)
        p.paragraph_format.space_after = Pt(2)
        p.add_run("%s   %s" % (guion, texto))
        return p

    def letra(self, letra, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("%s)  %s" % (letra, texto))
        r.bold = True
        return p

    def codigo(self, texto, tamanio=7.5):
        """Bloque de codigo monoespaciado dentro de un recuadro."""
        tabla = self.doc.add_table(rows=1, cols=1)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        celda = tabla.cell(0, 0)
        celda.width = Cm(16)
        sombrear(celda, "FBF7EF")
        celda.paragraphs[0]._p.getparent().remove(celda.paragraphs[0]._p)
        for linea in texto.split("\n"):
            p = celda.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(linea if linea.strip() else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(tamanio)
            r.font.color.rgb = NOCHE
            rPr = r._element.get_or_add_rPr()
            rf = rPr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rPr.append(rf)
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rf.set(qn(attr), "Consolas")
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return tabla

    def tabla(self, cabecera, filas, anchos):
        t = self.doc.add_table(rows=1, cols=len(cabecera))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, texto in enumerate(cabecera):
            celda = t.rows[0].cells[i]
            celda.width = Cm(anchos[i])
            sombrear(celda, "12B5B0")
            celda.paragraphs[0].text = ""
            r = celda.paragraphs[0].add_run(texto)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for n, fila in enumerate(filas):
            celdas = t.add_row().cells
            for i, texto in enumerate(fila):
                celdas[i].width = Cm(anchos[i])
                if n % 2 == 1:
                    sombrear(celdas[i], "F5F8FA")
                celdas[i].paragraphs[0].text = ""
                p = celdas[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(texto)
                r.font.size = Pt(9)
                if i == 0:
                    r.font.name = "Consolas"
                    rPr = r._element.get_or_add_rPr()
                    rf = OxmlElement("w:rFonts")
                    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                        rf.set(qn(attr), "Consolas")
                    rPr.append(rf)
                    r.font.size = Pt(8.5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    def salto(self):
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def guardar(self, ruta):
        self.doc.core_properties.title = "Memoria Descriptiva - SocialKids"
        self.doc.core_properties.author = "Proyecto SocialKids"
        self.doc.core_properties.subject = (
            "Aplicacion movil educativa para el desarrollo de la empatia, "
            "la comunicacion asertiva y las habilidades sociales"
        )
        self.doc.save(ruta)
        return ruta
