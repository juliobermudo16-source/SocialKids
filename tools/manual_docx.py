#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Genera el Manual de Usuario de SocialKids en formato .docx editable.

Sigue el modelo de manual proporcionado: portada, indice, introduccion y
secciones con pasos numerados junto a cada captura de pantalla.

Las capturas NO se incluyen: en su lugar se deja un marco vacio con el titulo
de la imagen, para que la persona que edita el documento pegue alli su propia
captura (clic dentro del marco -> Insertar -> Imagenes).

Uso:
    python tools/manual_docx.py [salida.docx]
"""
import os
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

TURQUESA = RGBColor(0x0E, 0x8F, 0x92)
CORAL = RGBColor(0xD9, 0x4A, 0x3A)
VIOLETA = RGBColor(0x6A, 0x4A, 0xE0)
SOL = RGBColor(0xC8, 0x8A, 0x14)
NOCHE = RGBColor(0x1B, 0x2B, 0x3F)
GRIS = RGBColor(0x55, 0x62, 0x72)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)


# ------------------------------------------------------------------ utiles

def sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def bordes_celda(celda, color="12B5B0", val="dashed", sz=12):
    tcPr = celda._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def sin_bordes(tabla):
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % lado)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def alto_fila(fila, cm):
    trPr = fila._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


def borde_inferior(parrafo, color="12B5B0", grosor=14):
    pPr = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(grosor))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    bordes.append(bottom)
    pPr.append(bordes)


# ------------------------------------------------------------------ manual

class Manual:

    def __init__(self):
        self.doc = Document()
        self.n_paso = 0
        self.n_imagen = 0
        self._configurar()

    def _configurar(self):
        s = self.doc.sections[0]
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = NOCHE
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        pie = s.footer.paragraphs[0]
        pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = pie.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = GRIS
        for tipo in ("begin", "PAGE", "end"):
            if tipo == "PAGE":
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = " PAGE "
            else:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), tipo)
            run._r.append(el)

    # --------------------------------------------------------------- bloques

    def portada(self):
        for _ in range(5):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("MANUAL PARA USUARIOS")
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = NOCHE

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run("SOCIALKIDS")
        r.bold = True
        r.font.size = Pt(46)
        r.font.color.rgb = TURQUESA

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Isla Conecta")
        r.font.size = Pt(16)
        r.font.color.rgb = CORAL

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        borde_inferior(p)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        r = p.add_run("EXPLORADOR")
        r.bold = True
        r.font.size = Pt(26)
        r.font.color.rgb = VIOLETA

        # Marco para el logotipo o una captura de portada
        self.marco("PORTADA DE LA APLICACION", alto=8.0, ancho=8.0, numerar=False)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Version 1.0.0")
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRIS

        self.salto()

    def indice(self, entradas, nota_edicion=None):
        p = self.doc.add_paragraph()
        r = p.add_run("INDICE")
        r.bold = True
        r.font.size = Pt(24)
        r.font.color.rgb = TURQUESA
        borde_inferior(p)
        self.doc.add_paragraph()

        t = self.doc.add_table(rows=0, cols=2)
        sin_bordes(t)
        for numero, titulo in entradas:
            fila = t.add_row().cells
            fila[0].width = Cm(2.4)
            fila[1].width = Cm(14.2)
            p0 = fila[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r0 = p0.add_run(numero)
            r0.bold = True
            r0.font.color.rgb = CORAL
            r0.font.size = Pt(11)
            p1 = fila[1].paragraphs[0]
            sangria = numero.count(".") - 1
            p1.paragraph_format.left_indent = Cm(0.4 + sangria * 0.5)
            r1 = p1.add_run(titulo)
            r1.font.size = Pt(11)
            r1.bold = (sangria == 0)
        if nota_edicion:
            self.doc.add_paragraph()
            self.aviso(nota_edicion[0], nota_edicion[1], "FDECEC", "D94A3A")
        self.salto()

    def titulo_seccion(self, texto, subtitulo=None, color="12B5B0"):
        t = self.doc.add_table(rows=1, cols=1)
        celda = t.cell(0, 0)
        celda.width = Cm(16.6)
        sombrear(celda, color)
        p = celda.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = BLANCO
        if subtitulo:
            p2 = celda.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            r2 = p2.add_run(subtitulo)
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = BLANCO
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def texto(self, contenido, cursiva=False, tamanio=11):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(contenido)
        r.italic = cursiva
        r.font.size = Pt(tamanio)
        return p

    def marco(self, titulo, alto=8.5, ancho=15.0, numerar=True):
        """Recuadro vacio donde pegar una captura de pantalla."""
        if numerar:
            self.n_imagen += 1
            etiqueta = "IMAGEN %d. %s" % (self.n_imagen, titulo.upper())
        else:
            etiqueta = titulo.upper()

        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        celda = t.cell(0, 0)
        celda.width = Cm(ancho)
        sombrear(celda, "F7FBFC")
        bordes_celda(celda)
        alto_fila(t.rows[0], alto)

        celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[ PEGA AQUI TU CAPTURA DE PANTALLA ]")
        r.font.size = Pt(10)
        r.font.color.rgb = GRIS
        r.italic = True

        pie = self.doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pie.paragraph_format.space_before = Pt(2)
        r = pie.add_run(etiqueta)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = CORAL
        return etiqueta

    def bloque(self, titulo_imagen, pasos, alto=8.5):
        """Una captura con sus pasos numerados debajo, como en el modelo."""
        self.marco(titulo_imagen, alto=alto)
        for texto in pasos:
            self.n_paso += 1
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.9)
            p.paragraph_format.first_line_indent = Cm(-0.9)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("%d. " % self.n_paso)
            r.bold = True
            r.font.color.rgb = TURQUESA
            p.add_run(texto)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def aviso(self, titulo, contenido, color="FFF4DC", borde="E0A93B"):
        t = self.doc.add_table(rows=1, cols=1)
        celda = t.cell(0, 0)
        celda.width = Cm(16.6)
        sombrear(celda, color)
        bordes_celda(celda, borde, "single", 8)
        p = celda.paragraphs[0]
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = SOL
        p2 = celda.add_paragraph()
        r2 = p2.add_run(contenido)
        r2.font.size = Pt(10)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def salto(self):
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def guardar(self, ruta):
        self.doc.core_properties.title = "Manual para Usuarios - SocialKids"
        self.doc.core_properties.author = "Proyecto SocialKids"
        self.doc.core_properties.subject = "Manual de usuario del explorador"
        self.doc.save(ruta)
        return ruta
